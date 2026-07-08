"""
Parses the incoming `message` field from the backend.

Two shapes are supported:

1. NEW (current real backend shape) — a plain string, sometimes containing a
   media marker instead of real text:
       "My name is Tareq, phone 01700000099"          <- normal text
       "[Media audio: http://localhost:5000/v1/whatsapp/media/<id>]"
       "[Media image: http://localhost:5000/v1/whatsapp/media/<id>]"
       "[Media document: https://yourdomain.com/uploads/.../file.pdf]"
       "[Media document: https://yourdomain.com/uploads/.../file.docx]"
       "[Media document: https://yourdomain.com/uploads/.../file.xlsx]"
       "[Media document: https://yourdomain.com/uploads/.../file.csv]"

   The document's actual format (PDF / DOCX / XLSX / CSV) is detected from
   the URL's own file extension — the bracket marker itself just says
   "document" for all of these.

   The media URL points at the BACKEND's own protected media proxy (NOT a
   public Meta/WhatsApp URL), so it must be fetched with the same
   `x-api-token` header used for every other backend call.

2. LEGACY (older / possibly still used elsewhere) — a list of typed blocks:
       [{"type": "text", "content": "..."}]
       [{"type": "image", "url": "https://..."}]
       [{"type": "document", "content": {"text": "..."}}]

   This shape is kept working for backward compatibility, unchanged from
   before.

Output of `parse_incoming_message`: a tuple `(llm_message_content, history_text)`
  - `llm_message_content`: what actually goes to the LLM this turn — a plain
    string, or a LangChain multimodal content list (when an image is present).
  - `history_text`: a plain-text version for long-term memory / the summary
    job. Image bytes/base64 are never stored long-term. Voice notes ARE
    stored as their transcribed text (with a 🎤 marker), and documents ARE
    stored as their extracted text, since both ARE the customer's message.

Returning both together (instead of two separate calls) avoids downloading /
transcribing the same media twice.
"""

import re
import io
import csv
import base64
from typing import Any, Union, Tuple
from urllib.parse import urlparse

import httpx
import fitz  # PyMuPDF — same library used in extract_and_ingest.py
from docx import Document as DocxDocument
from openpyxl import load_workbook
from openai import AsyncOpenAI

from core.config import OPENAI_API_KEY, ROBERTO_API_TOKEN

_openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)

# Matches: "[Media audio: <url>]", "[Media image: <url>]", or "[Media document: <url>]"
_MEDIA_PATTERN = re.compile(r"^\[Media (image|audio|document):\s*(.+)\]$")

_MEDIA_DOWNLOAD_TIMEOUT = 30.0


async def _download_media(url: str) -> bytes:
    """
    Downloads media from the backend's own protected media proxy.
    Requires x-api-token — this is NOT a public URL.
    """
    async with httpx.AsyncClient() as client:
        resp = await client.get(
            url,
            headers={"x-api-token": ROBERTO_API_TOKEN},
            timeout=_MEDIA_DOWNLOAD_TIMEOUT,
        )
        resp.raise_for_status()
        return resp.content


def _guess_image_mime(data: bytes) -> str:
    """Sniff common image formats from magic bytes. Defaults to jpeg."""
    if data.startswith(b"\xff\xd8"):
        return "image/jpeg"
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if data.startswith(b"GIF87a") or data.startswith(b"GIF89a"):
        return "image/gif"
    if len(data) >= 12 and data[0:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    return "image/jpeg"


def _guess_audio_filename(url: str) -> str:
    """
    Picks a filename (with extension) for the Whisper API based on the
    media URL's own extension. WhatsApp voice notes are usually .ogg, but
    other channels may send .mp3/.wav/.m4a — Whisper needs a matching
    extension to correctly parse the audio, so we don't hardcode one.
    """
    path = urlparse(url).path
    ext = path.rsplit(".", 1)[-1].lower() if "." in path.rsplit("/", 1)[-1] else ""
    valid_exts = {"ogg", "mp3", "wav", "m4a", "webm", "mp4", "mpeg", "mpga"}
    if ext in valid_exts:
        return f"voice_note.{ext}"
    return "voice_note.ogg"  # sensible default (WhatsApp's usual format)


async def _transcribe_audio(url: str) -> str:
    """
    Downloads a voice note and transcribes it via OpenAI Whisper.
    Returns the transcribed text, or a safe fallback placeholder on failure.
    """
    try:
        audio_bytes = await _download_media(url)
    except Exception as e:
        print(f"[VOICE ERROR] Failed to download audio from {url}: {e}")
        return "[Customer sent a voice message, but it could not be downloaded]"

    filename = _guess_audio_filename(url)

    try:
        transcript = await _openai_client.audio.transcriptions.create(
            model="whisper-1",
            file=(filename, audio_bytes, "audio/mpeg"),
        )
        text = (transcript.text or "").strip()
        print(f"[VOICE] Transcribed ({len(text)} chars): {text[:100]}")
        return text or "[Customer sent a voice message, but no speech was detected]"
    except Exception as e:
        print(f"[VOICE ERROR] Transcription failed: {e}")
        return "[Customer sent a voice message, but it could not be transcribed]"


def _get_extension(url: str) -> str:
    """Extracts the lowercase file extension from a URL's path."""
    path = urlparse(url).path
    name = path.rsplit("/", 1)[-1]
    if "." in name:
        return name.rsplit(".", 1)[-1].lower()
    return ""


def _extract_pdf_text(doc_bytes: bytes) -> str:
    doc = fitz.open(stream=doc_bytes, filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    doc.close()
    return full_text.strip()


def _extract_docx_text(doc_bytes: bytes) -> str:
    document = DocxDocument(io.BytesIO(doc_bytes))
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts).strip()


def _extract_xlsx_text(doc_bytes: bytes) -> str:
    workbook = load_workbook(io.BytesIO(doc_bytes), data_only=True)
    parts = []

    for sheet in workbook.worksheets:
        parts.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                row_text = " | ".join("" if cell is None else str(cell) for cell in row)
                parts.append(row_text)

    return "\n".join(parts).strip()


def _extract_csv_text(doc_bytes: bytes) -> str:
    """
    Decodes CSV bytes to text (trying common encodings) and reformats rows
    for readability, the same "row | row | row" style used for XLSX sheets.
    """
    text = None
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = doc_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = doc_bytes.decode("utf-8", errors="replace")

    parts = []
    reader = csv.reader(io.StringIO(text))
    for row in reader:
        if any(cell.strip() for cell in row):
            parts.append(" | ".join(row))

    return "\n".join(parts).strip()


async def _extract_document_text(url: str) -> str:
    """
    Downloads a document (PDF, DOCX, or XLSX — detected from the URL's
    extension) and extracts its text/content. Returns a customer-message
    -style string with the extracted content, or a safe fallback placeholder
    on failure.
    """
    try:
        doc_bytes = await _download_media(url)
    except Exception as e:
        print(f"[DOCUMENT ERROR] Failed to download document from {url}: {e}")
        return "[Customer shared a document, but it could not be downloaded]"

    ext = _get_extension(url)

    try:
        if ext == "pdf":
            full_text = _extract_pdf_text(doc_bytes)
        elif ext == "docx":
            full_text = _extract_docx_text(doc_bytes)
        elif ext in ("xlsx", "xlsm"):
            full_text = _extract_xlsx_text(doc_bytes)
        elif ext == "csv":
            full_text = _extract_csv_text(doc_bytes)
        elif ext == "doc":
            print(f"[DOCUMENT WARN] Legacy .doc not supported: {url}")
            return "[Customer shared an older .doc file, which isn't supported — ask them to resend as .docx or .pdf]"
        elif ext == "xls":
            print(f"[DOCUMENT WARN] Legacy .xls not supported: {url}")
            return "[Customer shared an older .xls file, which isn't supported — ask them to resend as .xlsx or .pdf]"
        else:
            print(f"[DOCUMENT WARN] Unsupported document extension '{ext}': {url}")
            return f"[Customer shared a document (.{ext or 'unknown'} file), but this file type isn't supported for reading]"

        print(f"[DOCUMENT] Extracted {len(full_text)} chars from .{ext} file")

        if not full_text:
            return "[Customer shared a document, but no readable text was found in it]"

        return f"[Customer shared a document with this content:]\n{full_text}"
    except Exception as e:
        print(f"[DOCUMENT ERROR] Text extraction failed (.{ext}): {e}")
        return "[Customer shared a document, but it could not be read]"


async def _image_to_data_url(url: str):
    """
    Downloads an image from the backend's protected media proxy and returns
    it as a base64 data: URL (since the LLM provider can't fetch a URL that
    requires our x-api-token header). Returns None on failure.
    """
    try:
        image_bytes = await _download_media(url)
    except Exception as e:
        print(f"[IMAGE ERROR] Failed to download image from {url}: {e}")
        return None

    mime = _guess_image_mime(image_bytes)
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime};base64,{b64}"


async def _parse_block_list(blocks: list) -> Tuple[Union[str, list], str]:
    """
    LEGACY handling for the old list-of-blocks message shape. Unchanged
    behaviour from before, kept for backward compatibility.
    """
    text_parts = []
    image_urls = []

    for block in blocks:
        if not isinstance(block, dict):
            continue

        block_type = block.get("type")

        if block_type == "text":
            content = block.get("content", "")
            if content:
                text_parts.append(str(content))

        elif block_type == "image":
            url = block.get("url")
            if url:
                image_urls.append(url)

        elif block_type == "document":
            doc_content = block.get("content", {})
            doc_text = doc_content.get("text", "") if isinstance(doc_content, dict) else ""
            if doc_text:
                text_parts.append(f"[Customer shared a document with this content:]\n{doc_text}")

        else:
            print(f"[MESSAGE PARSE WARNING] Unknown block type: {block_type}")

    combined_text = "\n".join(text_parts).strip()

    if not image_urls:
        plain = combined_text or "[Empty message]"
        return plain, plain

    content_list = []
    if combined_text:
        content_list.append({"type": "text", "text": combined_text})
    else:
        content_list.append({"type": "text", "text": "Customer shared an image."})

    for url in image_urls:
        content_list.append({"type": "image_url", "image_url": {"url": url}})

    history_text = combined_text
    history_text = (history_text + "\n[Customer shared an image]").strip() if history_text else "[Customer shared an image]"

    return content_list, history_text


async def parse_incoming_message(raw_message: Any) -> Tuple[Union[str, list], str]:
    """
    Normalizes `raw_message` into `(llm_message_content, history_text)`.

    - Plain string with no media marker -> unchanged, used as-is for both.
    - Plain string "[Media audio: <url>]" -> downloaded + transcribed via
      Whisper; the transcribed text is used for BOTH the LLM turn and history
      (the transcription IS the customer's message).
    - Plain string "[Media image: <url>]" -> downloaded + base64-encoded into
      a multimodal LLM content list; history stores a short placeholder
      instead (image bytes never go into long-term memory).
    - List of typed blocks (legacy shape) -> handled by _parse_block_list,
      unchanged from before.
    """
    if isinstance(raw_message, str):
        match = _MEDIA_PATTERN.match(raw_message.strip())

        if not match:
            # Ordinary text message — unchanged behaviour.
            return raw_message, raw_message

        media_type, url = match.group(1), match.group(2).strip()

        if media_type == "audio":
            transcribed = await _transcribe_audio(url)
            # LLM gets the raw transcribed text (natural response, no marker).
            # History/logs get a marker prefix so it's clear later this turn
            # came from a voice note, not typed text.
            history_with_marker = f"🎤 [Voice message]: {transcribed}"
            return transcribed, history_with_marker

        if media_type == "image":
            data_url = await _image_to_data_url(url)
            if data_url is None:
                fallback = "[Customer shared an image, but it could not be loaded]"
                return fallback, fallback

            content_list = [
                {"type": "text", "text": "Customer shared an image."},
                {"type": "image_url", "image_url": {"url": data_url}},
            ]
            return content_list, "[Customer shared an image]"

        if media_type == "document":
            doc_text = await _extract_document_text(url)
            # Both LLM content and history get the extracted text — same
            # pattern as the legacy block-list document handling.
            return doc_text, doc_text

        # Regex only matches image/audio/document, so this is unreachable —
        # fail safe.
        return raw_message, raw_message

    if isinstance(raw_message, list):
        return await _parse_block_list(raw_message)

    print(f"[MESSAGE PARSE WARNING] Unexpected message type: {type(raw_message)}")
    return "", "[Empty message]"